#!/usr/bin/env python3
"""Build report/ResQRules_Report.docx — the standalone committee report.

Follows the approved outline in PRESENTATION.md §1a. English; team nicknames
kept (to be replaced by the team before submission). Identity styling:
#CC0000 heading system, Noto Sans body (falls back if not installed),
Consolas for code and verbatim engine traces.
Regenerate:  python3 report/build_report.py
"""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = pathlib.Path(__file__).resolve().parents[1]
TRACES = ROOT / "presentation" / "traces"

PRIMARY   = RGBColor(0xCC, 0x00, 0x00)
PRIMARY_D = RGBColor(0x99, 0x00, 0x00)
INK       = RGBColor(0x1A, 0x1A, 0x1A)
INK2      = RGBColor(0x4A, 0x4A, 0x4A)
INK3      = RGBColor(0x6B, 0x6B, 0x6B)
SANS, MONO = "Noto Sans", "Consolas"

doc = Document()

# ── base styles ────────────────────────────────────────────────────────────
st = doc.styles["Normal"]
st.font.name = SANS; st.font.size = Pt(10.5); st.font.color.rgb = INK
st.element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
pf = st.paragraph_format
pf.space_after = Pt(7); pf.line_spacing = 1.22

for name, sz, col, before in [("Heading 1", 17, PRIMARY_D, 20),
                              ("Heading 2", 13.5, PRIMARY_D, 14),
                              ("Heading 3", 11.5, INK, 10)]:
    hs = doc.styles[name]
    hs.font.name = SANS; hs.font.size = Pt(sz); hs.font.bold = True
    hs.font.color.rgb = col
    hs.paragraph_format.space_before = Pt(before)
    hs.paragraph_format.space_after = Pt(5)

for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.4)

def shade(el, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    el.get_or_add_pr = None
    return shd

def P(text=None, runs=None, size=10.5, align=None, space_after=None, color=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    for t, kw in (runs or [(text, {})]):
        r = p.add_run(t)
        r.font.name = kw.get("font", SANS)
        r.font.size = Pt(kw.get("size", size))
        r.font.bold = kw.get("bold", False)
        r.font.italic = kw.get("italic", False)
        r.font.color.rgb = kw.get("color", color or INK)
    return p

def bullets(items, size=10.5):
    for runs in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for t, kw in runs:
            r = p.add_run(t)
            r.font.name = kw.get("font", SANS); r.font.size = Pt(size)
            r.font.bold = kw.get("bold", False); r.font.italic = kw.get("italic", False)
            r.font.color.rgb = kw.get("color", INK)

def code_block(text, size=8.5):
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.35)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        r.font.name = MONO; r.font.size = Pt(size); r.font.color.rgb = INK2
    P("", size=4, space_after=2)

def table(headers, rows, widths=None, size=9.5, header_fill="CC0000"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = doc.styles["Table Grid"]; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), header_fill)
        tcPr.append(shd)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h); r.font.name = SANS; r.font.size = Pt(size)
        r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            cell = t.cell(i, j)
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            mono = isinstance(v, tuple) and v[1] == "mono"
            txt = v[0] if isinstance(v, tuple) else v
            r = p.add_run(txt)
            r.font.name = MONO if mono else SANS; r.font.size = Pt(size)
            r.font.color.rgb = INK
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    P("", size=4, space_after=4)
    return t

def H1(n, text): doc.add_heading(f"{n}.  {text}" if n else text, level=1)
def H2(n, text): doc.add_heading(f"{n}  {text}", level=2)

def page_number_footer():
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r"); tEl = OxmlElement("w:t"); tEl.text = "1"
        run.append(tEl); fld.append(run)
        p._p.append(fld)
        for r_ in p.runs: r_.font.size = Pt(9); r_.font.color.rgb = INK3

# ═════════════════ COVER ═════════════════
P("", space_after=60)
P(runs=[("ResQRules", {"size": 42, "bold": True, "color": PRIMARY})],
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
P(runs=[("A Knowledge-Based System for Red Crescent First-Responder Protocols",
         {"size": 15, "bold": True, "color": INK2})], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
P(runs=[("Knowledge-Based Systems — Course Project Report", {"size": 12, "color": INK3})],
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
P(runs=[("Team", {"size": 11, "bold": True, "color": PRIMARY_D})], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(runs=[("M.GH  ·  M.SH  ·  M.JA  ·  M.Y.D  ·  M.Marwan", {"size": 11.5})],
  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
P(runs=[("Supervisor", {"size": 11, "bold": True, "color": PRIMARY_D})], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
P(runs=[("Miss Omama", {"size": 11.5})], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
P(runs=[("July 2026", {"size": 10.5, "color": INK3})], align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ═════════════════ TOC ═════════════════
doc.add_heading("Contents", level=1)
p = doc.add_paragraph()
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
run = OxmlElement("w:r"); tEl = OxmlElement("w:t")
tEl.text = "Right-click and choose “Update Field” to generate the table of contents."
run.append(tEl); fld.append(run); p._p.append(fld)
doc.add_page_break()

# ═════════════════ 1. INTRODUCTION ═════════════════
H1(1, "Introduction")
P("ResQRules is a knowledge-based system (KBS) that walks a first responder through the "
  "Syrian Arab Red Crescent (SARC) first-responder standard-operating-procedure flowcharts as an "
  "interactive, interruptible question-and-answer session. It is built on the Experta "
  "forward-chaining rules engine and is delivered as three cooperating layers: a terminal "
  "inference engine, a FastAPI web service, and a bilingual (English/Arabic) Flutter mobile "
  "application for use in the field.")
P("The problem it addresses is simple to state. The SARC SOP is a deck of twelve expert-validated "
  "paper flowcharts. On paper, they are static: under the adrenaline of a real emergency the "
  "responder must find the right chart, locate the right box, and follow the arrows correctly — "
  "while the patient's condition keeps evolving. The critical failure mode is that the patient "
  "does not follow the flowchart: a new life threat (catastrophic bleeding, respiratory arrest, "
  "loss of consciousness) can appear while the responder is standing in the middle of an entirely "
  "different protocol. Paper cannot re-route. A reasoning system can.")
P("Stakeholders. Field first responders get a guided, interruptible protocol companion; trainees "
  "get a safe environment to practise protocol decisions; and, for the Knowledge-Based Systems "
  "course, the project is a complete working case study of the concepts the course examines: "
  "knowledge/inference separation, working memory, forward chaining, conflict resolution through "
  "salience, and reasoning under uncertainty with certainty factors.")
P("Scope. Version 1 covers exactly five charts: #1 General Intervention Process & Patient "
  "Assessment (the root), #3 Cardio-Respiratory Arrest – Adult, #7 Heimlich Maneuver / Choking, "
  "#8 Upper Airway Management (modelled shallow, as a hub), and #10 Catastrophic Hemorrhage. "
  "The remaining charts are a planned Phase-2 extension that — by construction of the "
  "architecture — requires adding JSON knowledge files only, with no engine change.")

# ═════════════════ 2. DOMAIN BACKGROUND ═════════════════
H1(2, "Domain Background: How a Paramedic Thinks")
P("Chart #1 organises the entire intervention into four assessment phases, drawn as swim-lanes "
  "on the source chart. We use the chart's own vocabulary throughout (the charts say "
  "“Assessment”, never “survey”):")
table(["Phase", "Name on chart", "Content"],
      [["1", "Preliminary Assessment (SSS)", "Safety (staff, patient, bystanders) → Scene evaluation → Situation evaluation"],
       ["2", "Primary Assessment", "AVPU consciousness scale + ABCDE: Airway & axis, Breathing, Circulation, Disability, Expose & protect"],
       ["3", "Secondary Assessment", "History and vitals: SAMPLER, OPQRST, vital signs, GCS, PERRLA"],
       ["4", "Continuous Assessment", "Re-assessment and orientation to the most appropriate care centre"]],
      widths=[1.4, 5.2, 9.4])
P(runs=[("One structural exception. ", {"bold": True}),
        ("Between phases 1 and 2 the chart asks a single question — “Life threatening bleeding?” "
         "(node gen_04) — and, on “yes”, jumps directly to the Catastrophic Hemorrhage protocol "
         "before ABCDE ever runs. The chart therefore encodes clinical precedence positionally: "
         "circulation-threatening bleeding outranks the standard assessment order. This is a key "
         "observation for the design, because our inference engine encodes exactly the same "
         "precedence as rule salience (Section 6.2).", {})])
P(runs=[("A note on terminology. ", {"bold": True}),
        ("Popular first-aid mnemonics such as DRS or CAB do not appear anywhere in the SARC "
         "source charts, so this report deliberately does not use them. All stage names above "
         "were verified verbatim against the transcribed chart text.", {"italic": True})])

# ═════════════════ 3. KBS BACKGROUND ═════════════════
H1(3, "Knowledge-Based Systems Background")
P("A knowledge-based system separates what it knows from how it reasons. The knowledge base holds "
  "domain facts and structures; the inference engine applies generic reasoning to them. In a "
  "forward-chaining production system, the reasoning cycle is: facts are declared into working "
  "memory; a pattern-matching network (RETE) matches rule conditions against the facts; eligible "
  "rule activations enter an agenda; conflict resolution picks which activation fires; the fired "
  "rule modifies working memory; and the cycle repeats until quiescence.")
P("Experta is the Python implementation used in this course. The primitives that matter for this "
  "project: Fact subclasses define typed facts with fields; a KnowledgeEngine subclass hosts the "
  "rules; @Rule decorates a method with a condition over facts (with MATCH bindings, NOT, TEST); "
  "declare and retract modify working memory; and salience orders competing activations — the "
  "highest-salience eligible rule fires first, which is Experta's primary conflict-resolution "
  "mechanism.")
P("Experta 1.9.4 ships no native certainty-factor (CF) machinery (no FuzzyCLIPS-style support). "
  "Section 6.4 shows how we modelled certainty factors as ordinary data — a numeric field on a "
  "fact, thresholded by an ordinary rule — which keeps uncertainty inside the standard "
  "declare/match/fire cycle.")

# ═════════════════ 4. KNOWLEDGE ACQUISITION & REPRESENTATION ═════════════════
H1(4, "Knowledge Acquisition and Representation")
H2("4.1", "Sources")
P("Two sources were used, together: the official SARC SOP flowcharts (the authoritative document) "
  "and an expert in the loop — M.GH, a software engineer with paramedic experience on the team, "
  "who collected the source material, resolved ambiguities during extraction, and validated the "
  "transcribed knowledge. The project's rule zero: transcribe the chart, never “improve” medical "
  "content.")
H2("4.2", "Visual-only extraction")
P("Chart pages were rasterised to PNG at ~180 DPI (rasterize.py) and read as images. Text-layer "
  "extraction (pdftotext and similar) was explicitly forbidden: it scrambles box labels and drops "
  "the arrows — and the arrows are the logic. Every extracted node records the source page it "
  "was read from.")
H2("4.3", "The knowledge representation")
P("Each chart is one JSON file in data/ conforming to SCHEMA.md. A chart has meta information "
  "(chart id, bilingual title, source pages), an entry node id, and a dictionary of nodes:")
table(["Node type", "Meaning", "Continuation"],
      [["question", "a decision diamond", "options[]; each option's next → a local node id"],
       ["instruction", "a do-this step that flows on", "single next → a local node id"],
       ["action", "a terminal leaf (stop)", "none"],
       ["jump", "a one-way cross-chart / hub transfer", "target_chart (+ target_node), resolved in another file"]],
      widths=[2.6, 6.2, 7.2])
P("All display text is bilingual (parallel text_en / text_ar fields on nodes, answers, prompts and "
  "titles) while everything structural — node ids, edges, jump targets, thresholds, provenance — "
  "is language-independent. Routing therefore never depends on the display language; a missing "
  "translation falls back to English with a visible flag.")
H2("4.4", "Provenance: the faithfulness audit trail")
P("Every node carries a provenance tag: source (transcribed from a drawn box or arrow), derived "
  "(the literal meaning of a drawn box, such as a “repeat until…” loop-back edge), or added "
  "(tool-added for navigation only — never clinical content). The extracted knowledge base:")
table(["Chart", "File", "Nodes", "source", "derived", "added", "Arabic"],
      [["#1 General Assessment", ("general_assessment.json", "mono"), "11", "11", "0", "0", "11/11"],
       ["#3 CPR Adult", ("cpr_adult.json", "mono"), "19", "19", "0", "0", "19/19"],
       ["#7 Choking / Heimlich", ("choking.json", "mono"), "13", "12", "0", "1", "13/13"],
       ["#8 Upper Airway (shallow)", ("airway.json", "mono"), "8", "8", "0", "0", "8/8"],
       ["#10 Catastrophic Hemorrhage", ("hemorrhage.json", "mono"), "8", "8", "0", "0", "8/8"],
       ["Total", "", "59", "58", "0", "1", "59/59"]],
      widths=[4.6, 3.9, 1.5, 1.6, 1.6, 1.5, 1.6])
P("58 of the 59 nodes are direct transcriptions of drawn content; the single added node is a "
  "navigation-only safe-exit edge, flagged in the extraction report. No clinical content was "
  "invented anywhere.")
H2("4.5", "Validation")
P("validate.py checks every knowledge file and fails loudly (non-zero exit) on any violation: the "
  "entry node must exist; every option and next must resolve to a real local node or a jump "
  "target; no node may be unreachable; and every in-tree leaf must be a terminal action. "
  "Validation is language-independent and was run after every extraction.")

# ═════════════════ 5. ARCHITECTURE ═════════════════
H1(5, "Architecture")
H2("5.1", "The design we rejected")
P("The obvious implementation — one hand-written @Rule per flowchart box — was rejected at the "
  "start. Fifty-nine rules that each hard-code “if we are at box X and the answer is Y, go to "
  "box Z” trap the knowledge inside the logic: every medical revision becomes a code change, and "
  "the rules engine demonstrates no inference at all. It reduces Experta to a flowchart costume — "
  "exactly like a web application with static hard-coded pages and no database.")
H2("5.2", "The design we built")
P("ResQRules separates the two concerns completely. The knowledge is data: one JSON tree per "
  "chart, 59 nodes across 5 charts. The inference is generic: one Experta KnowledgeEngine with "
  "8 rules in total that can walk any tree conforming to the schema, plus three genuine "
  "inference layers on top (Section 6). This separation is the textbook definition of a "
  "knowledge-based system — and it is falsifiable: adding chart #6 to the system means dropping "
  "one JSON file into data/ with zero code change, since charts are auto-discovered.")
P("The full product stands on that separation: a FastAPI service exposes the engine over REST "
  "with a contract driven by node type from the knowledge files (not by internal rules), and a "
  "Flutter application renders whatever node the engine serves (Section 8). Changes to the rule "
  "base require no API change and no app change.")

# ═════════════════ 6. INFERENCE ENGINE ═════════════════
H1(6, "The Inference Engine")
P("engine.py (≈420 lines) hosts one KnowledgeEngine. Working memory is built from six Fact types:")
table(["Fact", "Field", "Role"],
      [[("Position", "mono"), "nid", "where the walk currently is"],
       [("Active", "mono"), "chart", "which protocol is active — the override re-entry guard"],
       [("Danger", "mono"), "kind", "a transient paramedic observation (keys b / n / p / u)"],
       [("Bleeding", "mono"), "tier", "the hemorrhage escalation tier — state that accumulates"],
       [("StillBleeding", "mono"), "answer", "a transient recheck answer (yes / no)"],
       [("SignOfLife", "mono"), "cf", "graded confidence there is a sign of life (certainty factor)"]],
      widths=[3.2, 1.8, 11.0])
P("The complete rule inventory — eight rules, ordered by salience:")
table(["Rule", "Salience", "Fires when / does"],
      [[("ov_bleeding", "mono"), "100", "Danger(catastrophic_bleeding): preempt, jump to hemorrhage entry"],
       [("ov_arrest", "mono"), "90", "Danger(not_breathing | no_pulse): preempt, jump to CPR-adult entry"],
       [("ov_unconscious", "mono"), "80", "Danger(unconscious): preempt, jump to CPR-adult entry"],
       [("hem_enter", "mono"), "50", "arriving at hem_01: declares Bleeding(tier=1), starts the escalation loop"],
       [("cf_sign_of_life", "mono"), "20", "SignOfLife(cf): thresholds the CF and routes the cpr_11 decision"],
       [("hem_apply", "mono"), "10", "Bleeding(tier=N) with no recheck answer: applies the tier-N intervention, prompts the recheck"],
       [("hem_decide", "mono"), "10", "Bleeding(tier=N) + StillBleeding: advance / hold / exit (Section 6.3)"],
       [("walk", "mono"), "0", "baseline generic walker: renders the current node from JSON, retracts/declares Position"]],
      widths=[3.2, 1.8, 11.0])
H2("6.1", "The generic walker (forward chaining)")
P("The walk rule matches Position and Active, renders the current node from the JSON knowledge, "
  "asks the responder, and answers by retracting the old Position and declaring the new one — at "
  "which point the RETE network re-matches and the next activation fires. That chain of "
  "declare → match → fire is the forward chaining. The walker carries NOT(Bleeding()) so it cedes "
  "control while the hemorrhage loop owns the patient, and it resolves jump nodes across charts "
  "(printing a graceful “would transfer to …” stub when the target chart is not loaded).")
H2("6.2", "Inference layer 1 — primary-survey overrides (salience and conflict resolution)")
P("At every prompt the responder may, instead of answering, assert a danger observation from a "
  "fixed menu: b catastrophic bleeding, n not breathing, p no pulse, u unconscious. Asserting one "
  "declares a Danger fact; a high-salience override rule preempts whatever node the walk is on "
  "and jumps to the correct protocol entry. Salience encodes clinical precedence — bleed-out "
  "kills fastest (100), then arrest (90), then unconsciousness (80), with the baseline walker at "
  "0. When two dangers are asserted simultaneously, Experta's conflict resolution fires the "
  "higher-salience override first; this is proven by an automated test, not hand-sequenced.")
P("Two guards make the mechanism safe: each override retracts the Danger it handled, so it fires "
  "exactly once per assertion; and if the target protocol is already the Active chart, the "
  "override re-prompts in place instead of re-entering (no restart loops). This asynchronous "
  "assertion is the capability a static flowchart fundamentally cannot offer.")
H2("6.3", "Inference layer 2 — hemorrhage escalation (working-memory accumulation)")
P("Chart #10's escalation is modelled semantically rather than as drawn back-edges. The tier is a "
  "fact — Bleeding(tier=N) — that accumulates across rule firings:")
table(["Tier", "Intervention (from the chart)", "Recheck", "On “still bleeding”"],
      [["1", "Direct compression, compression dressing, hemostatic dressing (hem_01)", "hem_02", "advance to tier 2"],
       ["2", "Second compression bandage on top of the first (hem_03)", "hem_04", "advance to tier 3"],
       ["3", "Tourniquet (hem_05) — terminal tier", "hem_06", "HOLD: no tier 4; route to urgent transport (hem_07)"]],
      widths=[1.3, 7.6, 1.9, 5.2])
P("hem_decide reasons over the current tier plus the re-asserted StillBleeding answer: still "
  "bleeding below the terminal tier retracts Bleeding(N) and declares Bleeding(N+1) — the engine "
  "never re-asks an exhausted tier, because the fact remembers. “Controlled” exits to the "
  "post-control transport. One deliberate, documented divergence from the drawn chart: the source "
  "draws an endless re-apply loop at the tourniquet (hem_06 “yes” → hem_05); the semantic model "
  "caps at the terminal tier and routes to urgent transport instead of spinning forever. The "
  "divergence is flagged, not hidden.")
H2("6.4", "Inference layer 3 — certainty factor (reasoning under uncertainty)")
P("Anchor: cpr_adult node cpr_11, “[1] Sign of life?”, justified by the chart's own annotations "
  "(“Gasping is not considered efficient breathing”; “Don't waste time checking pulse”) — an "
  "inherently uncertain vital-sign judgement. Since Experta has no native CF support, the CF is "
  "ordinary data: the prompt maps four answers to a numeric confidence, and one rule "
  "(cf_sign_of_life, salience 20) thresholds it and routes the decision:")
table(["Answer", "CF", "Routed to"],
      [["certain", "+1.0", "ventilation branch (cpr_12)"],
       ["likely", "+0.6", "ventilation branch (cpr_12)"],
       ["unsure", "+0.2", "CPR branch (cpr_14) — “if in doubt, compress”"],
       ["none", "−1.0", "CPR branch (cpr_14)"]],
      widths=[2.6, 1.8, 9.0])
P("Threshold = 0.5 (inclusive). The conservative bias lives in the scale, not the threshold: "
  "genuine doubt scores 0.2 — far below the cut — because the source treats gasping as "
  "not-breathing and a doubtful pulse as unreliable. The endpoints are unchanged from a plain "
  "yes/no; only the ambiguous middle re-routes, and it re-routes toward compressions. The CF is "
  "deliberately one notch softer than an override: a Danger assertion is a binary, high-salience "
  "hard jump for when the paramedic is sure; the CF tips a single decision when they are not.")
H2("6.5", "Mapping to the course rubric")
table(["KBS concept", "Where it lives in ResQRules"],
      [["Facts / working memory", "the six Fact types of the engine (table above)"],
       ["Forward chaining", "walk: RETE matches Position, fires, declares the next Position → chains"],
       ["Salience", "overrides at 100/90/80 vs. walker at 0 — salience encodes clinical precedence"],
       ["Conflict resolution", "two simultaneous dangers → the higher salience fires first (TEST 1)"],
       ["Working-memory accumulation", "Bleeding(tier=N) advances across firings (TESTS 3–5)"],
       ["Certainty factors", "SignOfLife(cf) thresholded by a rule; drives routing (TESTS 6–8)"],
       ["Hubs / sub-protocols", "jump nodes; cross-chart resolution; graceful stubs for absent charts"]],
      widths=[5.2, 10.8])

# ═════════════════ 7. WALKTHROUGH ═════════════════
H1(7, "Complete Flow Walkthrough")
P("The following trace is a real engine session (English UI), reproduced verbatim in Appendix D. "
  "It starts on the root chart and, at gen_04, answers “yes” to life-threatening bleeding — the "
  "knowledge-rich path that exercises a cross-chart jump and the escalation loop. Each step shows "
  "the working memory after the step, the rule that fired, and the change (Δ) it made:")
table(["Step", "Rule (sal)", "Working memory after", "Δ this step"],
      [["1  gen_01 Safety", "walk (0)", "Position(gen_01), Active(general_assessment)", "+ Position, + Active (initial)"],
       ["2  gen_02 Scene", "walk (0)", "Position(gen_02), Active(general_assessment)", "− Position(gen_01), + Position(gen_02)"],
       ["3  gen_03 → gen_04", "walk (0) ×2", "Position(gen_04), Active(general_assessment)", "− Position(gen_03), + Position(gen_04)"],
       ["4  gen_04 = yes", "walk (0)", "—", "the answer becomes the routing input"],
       ["5  jump to hemorrhage", "walk (0)", "Position(hem_01), Active(hemorrhage)", "− Active(general_assessment), + Active(hemorrhage)"],
       ["6  loop entered", "hem_enter (50)", "+ Bleeding(tier=1)", "+ Bleeding(tier=1) — state, not position"],
       ["7  tier 1 → still bleeding", "hem_decide (10)", "Bleeding(tier=2)", "− Bleeding(1), + Bleeding(2)"],
       ["8  tier 2 → still bleeding", "hem_decide (10)", "Bleeding(tier=3) — terminal", "− Bleeding(2), + Bleeding(3)"],
       ["9  tier 3 → controlled", "hem_decide → walk", "Position(hem_07)", "− Bleeding(3), − StillBleeding; leaf reached"]],
      widths=[3.4, 2.6, 5.4, 4.6], size=8.5)
P("Every step above was produced by the same eight generic rules — zero chart-specific code. The "
  "“no” branch at gen_04 (Primary Assessment → stable/unstable → Secondary → Continuous) is the "
  "plain walker at salience 0 throughout, and is also reproduced in Appendix D.")

# ═════════════════ 8. SYSTEM INTEGRATION ═════════════════
H1(8, "System Integration")
H2("8.1", "The FastAPI service")
P("The api/ package wraps the engine in a REST service. Its stability guarantee is an "
  "architectural consequence: response shapes are driven by node type from the JSON knowledge — "
  "never by internal rules, facts or salience. Adding or changing rules is a zero API change; a "
  "new chart is auto-discovered; authentication and persistence can be added later as middleware "
  "and storage without touching response shapes. Every text field in every response carries both "
  "languages, so the client can switch display language without a server round-trip.")
table(["Endpoint", "Purpose"],
      [[("GET /charts", "mono"), "list available charts (auto-discovered from data/)"],
       [("POST /sessions", "mono"), "start a session on a chart"],
       [("GET /sessions/{id}", "mono"), "reconnect / read current screen"],
       [("POST /sessions/{id}/step", "mono"), "answer the current screen (numbers, CF keys, or danger keys b/n/p/u)"],
       [("POST /sessions/{id}/reset", "mono"), "restart the session"],
       [("DELETE /sessions/{id}", "mono"), "end the session"]],
      widths=[5.0, 11.0])
P("Hand-raised errors return a machine-readable body with a stable code (for example "
  "session_ended), so clients branch on codes rather than matching message strings.")
H2("8.2", "The Flutter application")
P("The production mobile client is a separate repository (resq_rules) built with Clean "
  "Architecture (domain / data / presentation), Cubit for state management, and Retrofit/Dio for "
  "the API client. Its central design rule mirrors the server contract: rendering is switched on "
  "the server's node type (a ScreenType enum selects one widget per type), while the shell — "
  "app bar, loading bar, and the persistent danger panel — never changes. The danger panel is "
  "inference layer 1 as user interface: four always-visible zones (bleeding, not breathing, no "
  "pulse, unconscious) that send the override keys. Override feedback quotes the fired salience. "
  "The language toggle is fully client-side (both languages arrive in every response), and "
  "right-to-left mirroring for Arabic is handled by the platform, with Latin clinical acronyms "
  "(CPR, AED, SpO2) kept Latin.")
H2("8.3", "An honest note on the prototype")
P("An earlier, simpler Flutter prototype lives in this repository's mobile/ folder (raw HTTP + "
  "provider). It validated the API contract and was then rebuilt properly as the production app. "
  "Both are kept for transparency.")

# ═════════════════ 9. TESTING ═════════════════
H1(9, "Testing and Validation")
P("Eleven automated tests pin every inference behaviour to a KBS property. Engine tests "
  "(test_overrides.py) drive the real engine with scripted input; API tests (test_api.py) "
  "exercise the service layer:")
table(["Test", "What it proves"],
      [["Engine TEST 1", "conflict resolution: two simultaneous dangers → bleeding (sal 100) fires before arrest (90)"],
       ["Engine TEST 2", "override guard: asserting not_breathing while already in cpr_adult does not re-enter"],
       ["Engine TEST 3", "escalation: tiers strictly increase 1→2→3; terminal holds (no tier 4, no spin)"],
       ["Engine TEST 4", "early exits: “controlled” at tiers 1, 2 and 3 each exit cleanly"],
       ["Engine TEST 5", "integration: override into hemorrhage, then the loop runs as if reached by normal walk"],
       ["Engine TEST 6", "CF: certain/likely → ventilation branch"],
       ["Engine TEST 7", "CF: unsure/none → CPR branch (“if in doubt, compress”)"],
       ["Engine TEST 8", "CF boundary: cf = 0.5 is inclusive → ventilation; endpoints match plain yes/no"],
       ["API TEST 1", "pressing b on a hemorrhage screen is a guard (keeps escalating), not an exit"],
       ["API TEST 2", "arrest override out of hemorrhage jumps charts with clean loop state"],
       ["API TEST 3", "recheck semantics: “yes” escalates the tier; “no” exits to the terminal screen"]],
      widths=[3.0, 13.0])
P("In addition, validate.py gates every knowledge file (Section 4.5), and the walkthrough slides "
  "and this report's Appendix D use only verbatim transcripts of real engine sessions.")

# ═════════════════ 10. LIMITATIONS ═════════════════
H1(10, "Limitations and Future Work")
bullets([
 [("No return stack. ", {"bold": True}),
  ("Cross-chart jumps are one-way; a sub-protocol cannot yet return to where it was called from. "
   "A return-stack is the designed next inference increment.", {})],
 [("Unconscious routing simplification. ", {"bold": True}),
  ("Per the fixed danger-menu specification, “unconscious” routes to the CPR-adult entry (whose "
   "first step is the unresponsiveness/airway assessment). Clinically, an unconscious patient "
   "with pulse and breathing should route airway-first; a finer split is documented future work.", {})],
 [("One CF node. ", {"bold": True}),
  ("The certainty factor is applied where the source chart itself signals uncertainty (cpr_11). "
   "Extending CFs to other ambiguous decisions is possible with the same CF-as-data pattern.", {})],
 [("Phase-2 charts. ", {"bold": True}),
  ("Charts #2, #4, #5, #6, #9, #11, #12 plus two referenced-but-absent charts (Oxygen, Seizure) "
   "are JSON-only additions; today any jump to them degrades gracefully to a printed stub.", {})],
 [("Service hardening. ", {"bold": True}),
  ("Authentication and persistent storage are deliberately absent in v1; the API contract was "
   "designed so both are additive.", {})],
])

# ═════════════════ 11. CONCLUSION ═════════════════
H1(11, "Conclusion")
P("The central lesson of ResQRules is the oldest one in knowledge-based systems, demonstrated "
  "end-to-end: keep the knowledge as data and the inference generic. Because the five SARC "
  "charts live as validated, provenance-tracked JSON, eight generic Experta rules walk all 59 "
  "nodes — and the interesting engineering budget went where it belongs, into real inference: "
  "salience that encodes clinical precedence and resolves conflicting dangers, working memory "
  "that accumulates an escalation state the flowchart can only draw as a loop, and a certainty "
  "factor that turns an uncertain vital-sign judgement into a principled, conservative routing "
  "decision. The same separation paid for the product around the engine: a REST contract that "
  "rules cannot break, and a bilingual field application driven entirely by the knowledge. "
  "Paper cannot react; knowledge as data plus a small inference engine can.")

# ═════════════════ 12. ACKNOWLEDGMENTS ═════════════════
H1(12, "Acknowledgments")
P("We thank our supervisor, Miss Omama, for guidance that kept the knowledge honest and the "
  "inference real.")
table(["Member", "Contribution"],
      [["M.GH", "Software engineer with paramedic experience; knowledge source and reference collector; domain validation; Flutter development"],
       ["M.SH", "Python expert and AI engineer; the Experta inference engine"],
       ["M.JA", "Backend engineer; wrapped the engine in FastAPI"],
       ["M.Y.D", "Flutter developer and designer"],
       ["M.Marwan", "Flutter expert; UI designer; DevOps on the server"]],
      widths=[2.8, 13.2])

# ═════════════════ APPENDICES ═════════════════
doc.add_page_break()
H1(None, "Appendix A — Chart Inventory")
table(["#", "Chart", "Source file (repo)", "Modelling"],
      [["1", "General Intervention Process & Patient Assessment", "General Intervention Process & Pat Assessment.pdf", "root / spine"],
       ["3", "Cardio-Respiratory Arrest – Adult", "Non-Traumatic/Cardio-Respiratory Arrest Adult_validated.pdf", "full; CF node cpr_11"],
       ["7", "Heimlich Maneuver / Choking", "Non-Traumatic/Heimlich Maneuver_validated.pdf", "full; jumps into CPR"],
       ["8", "Upper Airway Management", "Non-Traumatic/Upper Airway Management_validated.pdf", "SHALLOW hub"],
       ["10", "Catastrophic Hemorrhage", "Traumatic/Catastrophic Hemorrhage_validated.pdf", "full; escalation loop"]],
      widths=[0.9, 5.6, 6.6, 2.9], size=8.5)

H1(None, "Appendix B — Schema Quick Reference")
code_block('''{
  "meta":  { "chart_id", "title_en", "title_ar?", "source", "source_pages", "notes?" },
  "entry": "<node_id>",
  "nodes": {
    "<node_id>": {
      "type": "question | instruction | action | jump",
      "text_en" (required) / "text_ar" (optional, falls back),
      "page": <n>,                      // source page — required
      "provenance": "source | derived | added",
      // question:    "options": [ { "answer_en", "answer_ar?", "next", "provenance?" } ]
      // instruction: "next"
      // jump:        "target_chart", "target_node"
      // action:      leaf — no continuation
      // CF node:     "cf": { "prompt_en/_ar", "scale": {certain,likely,unsure,none},
      //                      "threshold", "confident_next", "doubt_next" }
    } } }''', size=8)

H1(None, "Appendix C — API Summary")
P("Base URL (hosted): https://resqrules.duckdns.org — local development: uvicorn on "
  "http://localhost:8000. All requests and responses are JSON; every displayed text field is a "
  "bilingual pair {en, ar}. Screens carry a type discriminator (instruction, question, "
  "cf_question, hem_question, action, jump_stub) that fully determines the client rendering. "
  "Hand-raised errors carry a stable machine-readable code. See api/API.md for the full "
  "reference.")

H1(None, "Appendix D — Verbatim Engine Transcripts")
P("Captured from real sessions (LANG = \"en\"); the files live in presentation/traces/.")
for title, fname in [
    ("D.1 — Root chart, YES branch: jump + escalation to terminal exit", "walkthrough_yes_branch_en.txt"),
    ("D.2 — Root chart, NO branch: the plain walk to the leaf", "walkthrough_no_branch_en.txt"),
    ("D.3 — Danger override: “b” asserted at the first prompt (salience 100)", "override_bleeding_en.txt"),
    ("D.4 — Certainty factor: “unsure” at cpr_11 routes to compressions", "cf_unsure_en.txt")]:
    doc.add_heading(title, level=2)
    code_block((TRACES / fname).read_text(), size=7.5)

page_number_footer()
OUT = pathlib.Path(__file__).parent / "ResQRules_Report.docx"
doc.save(str(OUT))
print("saved", OUT)
